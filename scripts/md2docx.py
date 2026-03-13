import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    if len(sys.argv) < 3:
        print("Usage: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)
        
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    doc = Document()
    
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            _process_inline(p, text)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line)
            _process_inline(p, text)
        elif line.startswith('---'):
            doc.add_page_break()
        else:
            p = doc.add_paragraph()
            _process_inline(p, line)
            
    doc.save(output_file)
    print(f"Successfully converted {input_file} to {output_file}")

def _process_inline(paragraph, text):
    # Very basic processing for **bold** and `code`
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*'):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = 'Consolas'
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    main()
